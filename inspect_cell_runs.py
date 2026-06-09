import docx

def inspect_cell_runs(path):
    doc = docx.Document(path)
    table = doc.tables[0]
    cell = table.cell(0, 1)
    print("Paragraphs in cell R0C1:", len(cell.paragraphs))
    for p_idx, p in enumerate(cell.paragraphs):
        print(f"  P{p_idx}: text='{p.text}'")
        print(f"    Runs: {len(p.runs)}")
        for r_idx, run in enumerate(p.runs):
            print(f"      Run {r_idx}: text='{run.text}', font_name='{run.font.name}', font_size={run.font.size}")

if __name__ == "__main__":
    import os
    base_dir = os.path.dirname(os.path.abspath(__file__))
    inspect_cell_runs(os.path.join(base_dir, "GE Label - 42300276510.docx"))
