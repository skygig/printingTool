import docx
import os
import copy
from docx.table import Table

def generate_docx_labels(template_path, output_path, po_num=None, part_num=None, description=None, qty=4, items=None):
    """
    Modifies the grid-based part label DOCX template.
    Generates labels for each item according to its specified quantity,
    placing them sequentially within the label document.
    Preserves original layout, formatting, and styles using XML element cloning.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template docx file not found at: {template_path}")
        
    doc = docx.Document(template_path)
    
    # Check that tables exist
    if not doc.tables:
        raise ValueError("No tables found in the template docx.")
        
    body = doc.element.body
    
    # Store references to the 4 template table XML elements and a spacer paragraph XML element
    table_templates = [copy.deepcopy(t._tbl) for t in doc.tables]
    # In the template, all spacer paragraphs are empty. We'll use doc.paragraphs[0]._p as spacer.
    spacer_template = copy.deepcopy(doc.paragraphs[0]._p) if doc.paragraphs else None
    
    # Clear the body except for the section properties (sectPr)
    sectPr = body.find(docx.oxml.ns.qn('w:sectPr'))
    for child in list(body):
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)
        
    if items is None:
        items = [{
            'po_num': po_num,
            'part_num': part_num,
            'description': description,
            'qty': qty
        }]
        
    # Rebuild the document body with the labels for each item
    label_index = 0
    # Calculate total quantity across all items
    total_qty = 0
    valid_items = []
    for item in items:
        item_qty_str = item.get('qty', 0)
        try:
            item_qty = max(0, int(item_qty_str))
        except (ValueError, TypeError):
            item_qty = 0
        if item_qty > 0:
            total_qty += item_qty
            valid_items.append((item, item_qty))
            
    for item, item_qty in valid_items:
        item_po = item.get('po_num') or item.get('customer_po') or po_num
        item_part = item.get('part_num', '')
        item_desc = item.get('part_desc') or item.get('description') or ''
        
        for q in range(item_qty):
            # 1. Clone and insert table XML
            tbl_xml = copy.deepcopy(table_templates[label_index % 4])
            body.insert(body.index(sectPr), tbl_xml)
            
            # 2. Modify the cloned table contents (run-replacement to preserve styles)
            table = Table(tbl_xml, doc)
            if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
                cell = table.cell(0, 1)
                
                # Paragraph 2: GE PO# <po>
                if len(cell.paragraphs) > 2:
                    p2 = cell.paragraphs[2]
                    if len(p2.runs) >= 3:
                        p2.runs[2].text = str(item_po)
                    else:
                        p2.text = f"GE PO#   {item_po}"
                
                # Paragraph 3: Part# <part_num>
                if len(cell.paragraphs) > 3:
                    p3 = cell.paragraphs[3]
                    if len(p3.runs) >= 3:
                        p3.runs[2].text = str(item_part)
                    else:
                        p3.text = f"Part#       {item_part}"
                
                # Paragraph 4: Description: <description>
                if len(cell.paragraphs) > 4:
                    p4 = cell.paragraphs[4]
                    if len(p4.runs) >= 1:
                        p4.runs[0].text = f"Description:    {item_desc}"
                    else:
                        p4.text = f"Description:    {item_desc}                              "
            
            # 3. Add separator (spacers on same page, page break between pages)
            if label_index < total_qty - 1:
                if (label_index + 1) % 4 == 0:
                    # Add a page break paragraph before sectPr
                    p = doc.add_paragraph()
                    p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                    p_xml = p._p
                    body.remove(p_xml)
                    body.insert(body.index(sectPr), p_xml)
                else:
                    # Add two spacer paragraphs
                    if spacer_template is not None:
                        p1_xml = copy.deepcopy(spacer_template)
                        p2_xml = copy.deepcopy(spacer_template)
                        body.insert(body.index(sectPr), p1_xml)
                        body.insert(body.index(sectPr), p2_xml)
            label_index += 1
            
    # Save the modified document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Generated label document with {total_qty} labels at: {output_path}")
    return output_path

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    template = os.path.join(base_dir, "GE Label - 42300276510.docx")
    output = os.path.join(base_dir, "Outputs/test_label.docx")
    generate_docx_labels(template, output, "42300191168", "1-503-24-065", "TEFLON SEAL (TEST)", 6)
